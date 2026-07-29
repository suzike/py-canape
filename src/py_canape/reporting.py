"""审计、证据包、报告生成、脱敏和问题单闭环接口。"""

from __future__ import annotations

import hashlib
import json
import zipfile
from collections.abc import Iterable, Mapping, Sequence
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from html import escape
from pathlib import Path
from typing import Any, Protocol

from .errors import OptionalDependencyError


@dataclass(frozen=True, slots=True)
class AuditEntry:
    timestamp_utc: str
    action: str
    actor: str
    target: str
    status: str
    details: dict[str, Any]
    previous_hash: str
    hash: str


class AuditTrail:
    def __init__(self, path: str | Path) -> None:
        self.path = Path(path).expanduser().resolve()
        self.path.parent.mkdir(parents=True, exist_ok=True)

    def _last_hash(self) -> str:
        if not self.path.is_file():
            return "0" * 64
        lines = self.path.read_text(encoding="utf-8").splitlines()
        return json.loads(lines[-1])["hash"] if lines else "0" * 64

    def append(
        self,
        action: str,
        *,
        actor: str,
        target: str = "",
        status: str = "passed",
        details: Mapping[str, Any] | None = None,
    ) -> AuditEntry:
        previous = self._last_hash()
        payload = {
            "timestamp_utc": datetime.now(timezone.utc).isoformat(),
            "action": action,
            "actor": actor,
            "target": target,
            "status": status,
            "details": dict(details or {}),
            "previous_hash": previous,
        }
        canonical = json.dumps(payload, sort_keys=True, ensure_ascii=False, default=str)
        payload["hash"] = hashlib.sha256(canonical.encode("utf-8")).hexdigest()
        entry = AuditEntry(**payload)
        with self.path.open("a", encoding="utf-8") as stream:
            stream.write(json.dumps(asdict(entry), ensure_ascii=False, default=str) + "\n")
        return entry

    def verify(self) -> dict[str, Any]:
        previous = "0" * 64
        errors = []
        if not self.path.is_file():
            return {"passed": True, "entries": 0, "errors": []}
        lines = self.path.read_text(encoding="utf-8").splitlines()
        for number, line in enumerate(lines, start=1):
            item = json.loads(line)
            digest = item.pop("hash")
            canonical = json.dumps(
                item, sort_keys=True, ensure_ascii=False, default=str
            )
            expected = hashlib.sha256(canonical.encode("utf-8")).hexdigest()
            if digest != expected or item["previous_hash"] != previous:
                errors.append(number)
            previous = digest
        return {"passed": not errors, "entries": len(lines), "errors": errors}


class IssueTracker(Protocol):
    def attach(self, issue_id: str, path: str | Path) -> Any: ...

    def transition(self, issue_id: str, status: str) -> Any: ...


class Reporter:
    @staticmethod
    def anonymize(
        records: Sequence[Mapping[str, Any]],
        *,
        remove_fields: Iterable[str] = (),
        hash_fields: Iterable[str] = (),
        salt: str = "",
    ) -> list[dict[str, Any]]:
        remove = set(remove_fields)
        hashed = set(hash_fields)
        output = []
        for record in records:
            item = {}
            for key, value in record.items():
                if key in remove:
                    continue
                if key in hashed:
                    value = hashlib.sha256(f"{salt}{value}".encode()).hexdigest()
                item[key] = value
            output.append(item)
        return output

    @staticmethod
    def create_evidence_bundle(
        output_file: str | Path,
        *,
        files: Iterable[str | Path] = (),
        metadata: Mapping[str, Any] | None = None,
    ) -> Path:
        output = Path(output_file).expanduser().resolve()
        output.parent.mkdir(parents=True, exist_ok=True)
        manifest = {"created_utc": datetime.now(timezone.utc).isoformat(), "files": []}
        with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for value in files:
                path = Path(value).expanduser().resolve()
                if not path.is_file():
                    raise FileNotFoundError(path)
                digest = hashlib.sha256(path.read_bytes()).hexdigest()
                archive.write(path, arcname=f"evidence/{path.name}")
                manifest["files"].append(
                    {"name": path.name, "size": path.stat().st_size, "sha256": digest}
                )
            archive.writestr(
                "manifest.json",
                json.dumps(
                    {**manifest, "metadata": dict(metadata or {})},
                    ensure_ascii=False,
                    indent=2,
                    default=str,
                ),
            )
        return output

    @staticmethod
    def html(
        output_file: str | Path,
        *,
        title: str,
        sections: Sequence[tuple[str, Any]],
    ) -> Path:
        output = Path(output_file).expanduser().resolve()
        output.parent.mkdir(parents=True, exist_ok=True)
        parts = [
            "<!doctype html><html><head><meta charset='utf-8'>",
            f"<title>{escape(title)}</title>",
            "<style>body{font-family:Segoe UI,Arial;max-width:1100px;margin:2rem auto}"
            "pre{background:#f4f5f7;padding:1rem;overflow:auto}"
            "table{border-collapse:collapse}td,th{border:1px solid #ccc;padding:.4rem}</style>",
            "</head><body>",
            f"<h1>{escape(title)}</h1>",
        ]
        for heading, content in sections:
            parts.append(f"<h2>{escape(heading)}</h2>")
            serialized = json.dumps(content, ensure_ascii=False, indent=2, default=str)
            parts.append(
                f"<pre>{escape(serialized)}</pre>"
            )
        parts.append("</body></html>")
        output.write_text("".join(parts), encoding="utf-8")
        return output

    @staticmethod
    def excel(
        output_file: str | Path,
        sheets: Mapping[str, Sequence[Mapping[str, Any]]],
    ) -> Path:
        try:
            import openpyxl
        except ImportError as exc:
            raise OptionalDependencyError(
                "Excel 报告需要安装 py-canape-local[reports]"
            ) from exc
        output = Path(output_file).expanduser().resolve()
        output.parent.mkdir(parents=True, exist_ok=True)
        workbook = openpyxl.Workbook()
        workbook.remove(workbook.active)
        for title, rows_value in sheets.items():
            rows = list(rows_value)
            sheet = workbook.create_sheet(title=title[:31])
            headers = sorted({key for row in rows for key in row})
            if headers:
                sheet.append(headers)
                for row in rows:
                    sheet.append([row.get(header) for header in headers])
                sheet.freeze_panes = "A2"
        workbook.save(output)
        return output

    @staticmethod
    def word(
        output_file: str | Path,
        *,
        title: str,
        sections: Sequence[tuple[str, Any]],
    ) -> Path:
        try:
            from docx import Document
        except ImportError as exc:
            raise OptionalDependencyError(
                "Word 报告需要安装 py-canape-local[reports]"
            ) from exc
        output = Path(output_file).expanduser().resolve()
        output.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        document.add_heading(title, 0)
        for heading, content in sections:
            document.add_heading(heading, level=1)
            document.add_paragraph(
                json.dumps(content, ensure_ascii=False, indent=2, default=str)
            )
        document.save(output)
        return output

    @staticmethod
    def pdf(
        output_file: str | Path,
        *,
        title: str,
        sections: Sequence[tuple[str, Any]],
    ) -> Path:
        """生成无第三方依赖的基础 PDF；复杂企业版式可由报告插件覆盖。"""
        output = Path(output_file).expanduser().resolve()
        output.parent.mkdir(parents=True, exist_ok=True)
        text = [title]
        for heading, content in sections:
            text.extend(
                [
                    heading,
                    json.dumps(content, ensure_ascii=True, default=str),
                ]
            )
        lines = []
        for item in text:
            escaped = item.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            lines.extend(escaped.splitlines())
        commands = ["BT", "/F1 10 Tf", "50 790 Td"]
        for index, line in enumerate(lines[:48]):
            if index:
                commands.append("0 -15 Td")
            commands.append(f"({line[:100]}) Tj")
        commands.append("ET")
        stream = "\n".join(commands).encode("ascii", errors="replace")
        objects = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
            f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        ]
        content = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for number, body in enumerate(objects, start=1):
            offsets.append(len(content))
            content.extend(f"{number} 0 obj\n".encode())
            content.extend(body)
            content.extend(b"\nendobj\n")
        xref = len(content)
        content.extend(f"xref\n0 {len(objects) + 1}\n".encode())
        content.extend(b"0000000000 65535 f \n")
        for offset in offsets[1:]:
            content.extend(f"{offset:010d} 00000 n \n".encode())
        content.extend(
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref}\n%%EOF\n".encode()
        )
        output.write_bytes(content)
        return output

    @classmethod
    def generate(
        cls,
        output_file: str | Path,
        *,
        title: str,
        sections: Sequence[tuple[str, Any]],
    ) -> Path:
        suffix = Path(output_file).suffix.casefold()
        if suffix == ".html":
            return cls.html(output_file, title=title, sections=sections)
        if suffix == ".docx":
            return cls.word(output_file, title=title, sections=sections)
        if suffix == ".pdf":
            return cls.pdf(output_file, title=title, sections=sections)
        if suffix == ".xlsx":
            sheets = {
                heading: (
                    content
                    if isinstance(content, list)
                    else [{"value": json.dumps(content, ensure_ascii=False, default=str)}]
                )
                for heading, content in sections
            }
            return cls.excel(output_file, sheets)
        raise ValueError(f"不支持的报告格式：{suffix}")

    @staticmethod
    def publish_issue(
        tracker: IssueTracker,
        issue_id: str,
        evidence: str | Path,
        *,
        status: str | None = None,
    ) -> dict[str, Any]:
        attachment = tracker.attach(issue_id, evidence)
        transition = tracker.transition(issue_id, status) if status else None
        return {"attachment": attachment, "transition": transition}
